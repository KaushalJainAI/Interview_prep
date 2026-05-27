"""Convert a markdown file to .docx with code blocks, tables, and embedded PNGs.

Lightweight converter tuned for this notes project. Supports:
- # / ## / ### headings
- paragraphs (with **bold**, *italic*, `code`)
- fenced ``` code blocks (with language hint)
- GFM-style pipe tables
- bullet lists (-, *)
- numbered lists (1.)
- horizontal rules (---)
- > blockquotes
- image references via   `![alt](diagrams/foo.png)`  OR  inline HTML <img src="...">
"""
from __future__ import annotations
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
IMG  = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HTML_IMG = re.compile(r'<img\s+[^>]*src="([^"]+)"', re.I)

def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_runs(paragraph, text):
    """Apply inline formatting: bold/italic/code/links."""
    # protect images already handled
    pos = 0
    # We'll tokenize sequentially handling bold/italic/code/link in priority
    tokens = []
    i = 0
    while i < len(text):
        # try patterns at this position
        m_bold = BOLD.match(text, i)
        m_code = INLINE_CODE.match(text, i)
        m_link = LINK.match(text, i)
        m_ital = ITALIC.match(text, i)
        candidates = [m for m in (m_bold, m_code, m_link, m_ital) if m]
        if not candidates:
            tokens.append(("plain", text[i]))
            i += 1
            continue
        m = min(candidates, key=lambda mm: mm.start())
        if m.start() > i:
            tokens.append(("plain", text[i:m.start()]))
        if m is m_bold:
            tokens.append(("bold", m.group(1)))
        elif m is m_code:
            tokens.append(("code", m.group(1)))
        elif m is m_link:
            tokens.append(("link", (m.group(1), m.group(2))))
        elif m is m_ital:
            tokens.append(("ital", m.group(1)))
        i = m.end()

    # collapse adjacent plain
    out = []
    for kind, val in tokens:
        if kind == "plain" and out and out[-1][0] == "plain":
            out[-1] = ("plain", out[-1][1] + val)
        else:
            out.append((kind, val))

    for kind, val in out:
        if kind == "plain":
            r = paragraph.add_run(val)
        elif kind == "bold":
            r = paragraph.add_run(val); r.bold = True
        elif kind == "ital":
            r = paragraph.add_run(val); r.italic = True
        elif kind == "code":
            r = paragraph.add_run(val); r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x0b, 0x46, 0x8c)
        elif kind == "link":
            txt, _url = val
            r = paragraph.add_run(txt); r.font.color.rgb = RGBColor(0x1f, 0x6f, 0xeb); r.underline = True

def add_code_block(doc, code, lang=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # shaded background via paragraph properties
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), 'F4F6F8')
    pPr.append(shd)
    # border
    pBdr = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),'DDDDDD')
        pBdr.append(b)
    pPr.append(pBdr)
    for i, line in enumerate(code.rstrip("\n").split("\n")):
        if i: p.add_run().add_break()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def add_table(doc, rows):
    if not rows: return
    ncols = max(len(r) for r in rows)
    rows = [r + [""]*(ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_shading(cell, "E6F0FF")
    # set column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(6.5 / ncols)

def add_image(doc, src, base_dir):
    img_path = (base_dir / src).resolve()
    if not img_path.exists():
        doc.add_paragraph(f"[missing image: {src}]")
        return
    doc.add_picture(str(img_path), width=Inches(6.3))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def parse_table(lines):
    """Given a chunk of lines starting with '|', return list of cell rows."""
    rows = []
    for ln in lines:
        if not ln.strip().startswith("|"): break
        if re.match(r"^\s*\|?\s*:?-+:?\s*\|", ln):  # separator row
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows

def convert(md_path: Path, docx_path: Path, base_dir: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    doc = Document()

    # default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

    i = 0
    n = len(lines)
    while i < n:
        ln = lines[i]
        stripped = ln.rstrip()

        # code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            buf = []
            while i < n and not lines[i].rstrip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1  # skip closing fence
            add_code_block(doc, "\n".join(buf), lang)
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            head_text = m.group(2)
            p = doc.add_heading("", level=min(level, 4))
            add_runs(p, head_text)
            i += 1
            continue

        # image-only line
        m = IMG.match(stripped) or (HTML_IMG.search(stripped) and HTML_IMG.search(stripped))
        if IMG.match(stripped):
            mm = IMG.match(stripped)
            add_image(doc, mm.group(2), base_dir)
            i += 1; continue
        mm = HTML_IMG.search(stripped)
        if mm and stripped.startswith("<"):
            add_image(doc, mm.group(1), base_dir)
            i += 1; continue

        # horizontal rule
        if re.match(r"^-{3,}\s*$", stripped) or re.match(r"^\*{3,}\s*$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            b = OxmlElement('w:bottom')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:color'),'888888')
            pBdr.append(b); pPr.append(pBdr)
            i += 1; continue

        # table
        if stripped.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = parse_table(block)
            add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, stripped[2:])
            r = p.runs[0] if p.runs else None
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55,0x55,0x55)
            i += 1; continue

        # bullet
        m = re.match(r"^\s*[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1; continue

        # numbered
        m = re.match(r"^\s*\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1; continue

        # blank line
        if stripped == "":
            i += 1; continue

        # paragraph (possibly multi-line until blank)
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|```|\||> |- |\* |\d+\.\s|!\[)", lines[i].strip()
        ):
            para_lines.append(lines[i].strip()); i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(para_lines))

    doc.save(str(docx_path))

if __name__ == "__main__":
    md = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    base = md.parent
    convert(md, out, base)
    print("wrote", out)
